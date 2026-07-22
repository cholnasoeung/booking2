# -*- coding: utf-8 -*-
import docx
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HEADER_FILL = "1F497D"
STRIPE_A = "F0F4F8"
STRIPE_B = "FFFFFF"
COL_W = "2139"

ROWS = [
    ["User", "_id", "—", "Stores passenger and admin accounts; passwords hashed with bcryptjs"],
    ["Route", "_id", "—", "Defines origin, destination, and stops; referenced by Bus"],
    ["Bus", "_id", "routeId (Route), driverId (Driver)", "Each document represents one scheduled trip; embeds seat array"],
    ["BusDetail", "_id", "—", "Vehicle profile (registration, seat layout, amenities, documents); referenced by Bus, DriverSchedule, FuelLog, Maintenance, Incident"],
    ["RouteStop", "_id", "routeId (Route)", "Intermediate pickup/drop points along a route"],
    ["Booking", "_id", "userId (User), busId (Bus)", "One booking per transaction; embeds passenger details; TTL index on pending status"],
    ["PendingBooking", "_id", "userId, busId (string refs)", "Temporary checkout hold before payment confirms; TTL auto-expire"],
    ["WaitingList", "_id", "user (User), bus (Bus), route (Route)", "Queue of passengers waiting for a sold-out bus"],
    ["Driver", "_id", "—", "Stores driver profiles; assigned to bus schedules"],
    ["Employee", "_id", "—", "Non-driver staff profiles (mechanics, agents, managers, etc.)"],
    ["Attendance", "_id", "staffId (Driver/Employee)", "Daily attendance record; unique per staff + date"],
    ["DriverSchedule", "_id", "driverId (Driver), busDetailId (BusDetail), busId (Bus)", "Assigns a driver and vehicle to a shift"],
    ["DriverEarning", "_id", "driverId (Driver), busDetailId (BusDetail)", "Daily earnings computed from trips"],
    ["LeaveRequest", "_id", "staffId (Driver/Employee), reviewedBy (User)", "Leave applications and approval workflow"],
    ["Payroll", "_id", "employeeId (Employee)", "Monthly payroll run; unique per employee + month"],
    ["FuelLog", "_id", "busDetailId (BusDetail), driverId (Driver)", "Refueling records for cost tracking"],
    ["Maintenance", "_id", "busDetailId (BusDetail)", "Service and repair history with next-service reminders"],
    ["Incident", "_id", "busDetailId (BusDetail)", "Breakdown and accident reports"],
    ["PromoCode", "_id", "applicableRoutes (Route)", "Discount codes with usage limits and validity windows"],
    ["PricingRule", "_id", "routeId (Route)", "Seasonal and dynamic price multipliers"],
    ["Rating", "_id", "user (User), bus (Bus), booking (Booking)", "Passenger reviews and star ratings; unique per user + booking"],
    ["Loyalty", "_id", "user (User)", "Points balance, tier, and redemption history; one per user"],
    ["LostFound", "_id", "reportedBy (User), bookingId (Booking), busId (Bus), routeId (Route)", "Lost item case tracking"],
    ["Notification", "_id", "userId (User), busId (Bus), bookingId (Booking)", "In-app notifications; TTL auto-delete after 90 days"],
    ["SupportConversation", "_id", "user (User)", "Embedded support chat thread between passenger and admin"],
    ["AuditLog", "_id", "—", "Immutable action log for compliance; TTL auto-delete after 1 year"],
    ["PageView", "_id", "—", "Anonymous analytics event per page visit"],
    ["Settings", "_id", "—", "Singleton document storing site-wide configuration"],
]
assert len(ROWS) == 28

d = docx.Document("BUS_BOOKING_UPDATED.docx")
body = d.element.body
children = list(body)

def el_kind(el):
    if el.tag == qn('w:p'):
        return 'p'
    if el.tag == qn('w:tbl'):
        return 'tbl'
    return el.tag

def para_text(p_el):
    from docx.text.paragraph import Paragraph
    return Paragraph(p_el, d).text

start = None
for i, el in enumerate(children):
    if el_kind(el) == 'p' and para_text(el).strip() == "2.7 Database Schema Design":
        start = i
        break
assert start is not None

# Find the table and intro paragraph within the next few elements
intro_idx = None
tbl_idx = None
for i in range(start, start + 6):
    el = children[i]
    if el_kind(el) == 'p' and "five primary collections" in para_text(el):
        intro_idx = i
    if el_kind(el) == 'tbl':
        tbl_idx = i
        break
assert intro_idx is not None and tbl_idx is not None
print("intro_idx", intro_idx, "tbl_idx", tbl_idx)

# --- 1. Update intro paragraph text ---
from docx.text.paragraph import Paragraph
intro_p = Paragraph(children[intro_idx], d)
old_text = intro_p.text
new_text = (
    "The system uses MongoDB as a NoSQL document store, organized into 28 collections across "
    "seven functional domains: Transport, Booking, HR, Operations, Commerce, Communication, and "
    "System. Table 2.31 summarizes every collection and its relationships."
)
# clear existing runs, add one run with new text (keep paragraph formatting/pPr)
for r in list(intro_p.runs):
    r._r.getparent().remove(r._r)
run = intro_p.add_run(new_text)
run.font.name = "Times New Roman"
run.font.size = Pt(12)
print("old intro:", repr(old_text))
print("new intro:", repr(new_text))

# --- 2. Rebuild the table with all 28 rows, matching existing style ---
old_tbl_el = children[tbl_idx]

def set_shd(tcPr, fill):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def make_tcPr(fill):
    tcPr = OxmlElement('w:tcPr')
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:type'), 'dxa')
    tcW.set(qn('w:w'), COL_W)
    tcPr.append(tcW)
    set_shd(tcPr, fill)
    return tcPr

def make_run(text, bold=False, color=None, size_pt=10):
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Times New Roman')
    rPr.append(rFonts)
    if bold:
        rPr.append(OxmlElement('w:b'))
    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(size_pt * 2))
    rPr.append(sz)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    return r

def make_row(cells_text, fill, bold=False, color=None):
    tr = OxmlElement('w:tr')
    for text in cells_text:
        tc = OxmlElement('w:tc')
        tc.append(make_tcPr(fill))
        p = OxmlElement('w:p')
        p.append(make_run(text, bold=bold, color=color))
        tc.append(p)
        tr.append(tc)
    return tr

new_tbl = OxmlElement('w:tbl')
tblPr = OxmlElement('w:tblPr')
tblStyle = OxmlElement('w:tblStyle')
tblStyle.set(qn('w:val'), 'TableGrid')
tblPr.append(tblStyle)
tblW = OxmlElement('w:tblW')
tblW.set(qn('w:type'), 'auto')
tblW.set(qn('w:w'), '0')
tblPr.append(tblW)
tblLook = OxmlElement('w:tblLook')
tblLook.set(qn('w:val'), '04A0')
tblLook.set(qn('w:firstRow'), '1')
tblLook.set(qn('w:lastRow'), '0')
tblLook.set(qn('w:firstColumn'), '1')
tblLook.set(qn('w:lastColumn'), '0')
tblLook.set(qn('w:noHBand'), '0')
tblLook.set(qn('w:noVBand'), '1')
tblPr.append(tblLook)
new_tbl.append(tblPr)

tblGrid = OxmlElement('w:tblGrid')
for _ in range(4):
    gc = OxmlElement('w:gridCol')
    gc.set(qn('w:w'), COL_W)
    tblGrid.append(gc)
new_tbl.append(tblGrid)

new_tbl.append(make_row(["Collection", "Primary Key", "Foreign Keys", "Notes"],
                         HEADER_FILL, bold=True, color="FFFFFF"))
for i, row in enumerate(ROWS):
    fill = STRIPE_A if (i % 2 == 0) else STRIPE_B
    new_tbl.append(make_row(row, fill))

old_tbl_el.addprevious(new_tbl)
old_tbl_el.getparent().remove(old_tbl_el)

# --- 3. Fix inline "Table 2.7 summarizes" reference to "Table 2.31" ---
fixed_inline = 0
for p in d.paragraphs:
    for r in p.runs:
        if "Table 2.7 summarizes" in r.text:
            r.text = r.text.replace("Table 2.7 summarizes", "Table 2.31 summarizes")
            fixed_inline += 1
print("fixed inline refs:", fixed_inline)

d.save("BUS_BOOKING_UPDATED.docx")
print("Saved.")
